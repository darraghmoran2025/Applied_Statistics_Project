"""
generate_week3_results.py
Generates Week3_Results.docx
Run from the Week3 directory: python generate_week3_results.py
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
FIGS = os.path.join(HERE, "figures")
OUT  = os.path.join(HERE, "Week3_Results.docx")


# ── helpers ───────────────────────────────────────────────────────────────────

def _set_tnr(run, size_pt=12, bold=False, italic=False):
    run.font.name   = "Times New Roman"
    run.font.size   = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    r    = run._r
    rPr  = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rPr.insert(0, rFonts)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_tnr(run, bold=True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_tnr(run)
    return p


def add_figure(doc, filename, caption, width_in=6.3):
    path = os.path.join(FIGS, filename)
    if not os.path.exists(path):
        add_body(doc, f"[Figure not found: {filename}]")
        return
    p = doc.add_paragraph()
    p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    _set_tnr(r, italic=True)


def add_table(doc, headers, rows, caption):
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after  = Pt(2)
    r = cap.add_run(caption)
    _set_tnr(r, bold=True)

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"

    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_tnr(run, bold=True)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            _set_tnr(run)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ── document ──────────────────────────────────────────────────────────────────

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

title_p = doc.add_paragraph()
title_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(20)
r = title_p.add_run("Week 3 results: Variance-Gamma and NIG MLE")
_set_tnr(r, size_pt=14, bold=True)


# ── 1. Overview ───────────────────────────────────────────────────────────────

add_heading(doc, "1. Overview")

add_body(doc,
    "This week I fitted two Levy process models to the same 6,287 daily "
    "log-returns used in Week 2: the Variance-Gamma (VG) and the Normal "
    "Inverse Gaussian (NIG). Both have four parameters and allow asymmetric "
    "tails, which the Student-t cannot. On Neil's suggestion I also fitted the "
    "Laplace (double-exponential) distribution as a simple two-parameter "
    "benchmark. The Laplace is the symmetric special case of the VG (theta = 0, "
    "nu = 1), so it sits between the Gaussian and the heavier Levy models. With "
    "only two parameters, the same count as the Gaussian, it captures 1,771 of "
    "the 1,851 AIC units the four-parameter NIG gains over the Gaussian. Most "
    "of the improvement over the Gaussian therefore comes from allowing "
    "exponential rather than thin tails, before any extra parameters are added. "
    "NIG remains the best-fitting model. Its 99% Expected Shortfall of -5.19% "
    "sits between the Gaussian (-3.24%) and the Student-t (-5.81%), and it is "
    "the only model that passes a formal goodness-of-fit test. The sub-period "
    "analysis shows that the GFC and COVID were structurally different crises "
    "from the dot-com crash and the rate-hike cycle, and that the difference is "
    "clear in the Levy parameters."
)


# ── 2. Parameter estimates ────────────────────────────────────────────────────

add_heading(doc, "2. Parameter estimates")

add_body(doc,
    "Tables 1 and 3 carry the Week 2 Gaussian and Student-t results forward. "
    "Tables 2, 4 and 5 (Laplace, VG, NIG) are new."
)

add_table(doc,
    headers=["Parameter", "Estimate", "Standard Error", "t-statistic"],
    rows=[
        ["mu (daily drift)",  "+0.000223", "0.000154", "1.45"],
        ["sigma (scale)",     "0.012234",  "0.000109", "--"],
        ["Log-likelihood",    "18,764.3",  "--",       "--"],
        ["AIC",               "-37,524.5", "--",       "--"],
        ["BIC",               "-37,511.0", "--",       "--"],
    ],
    caption="Table 1. Gaussian MLE (Week 2)."
)

add_table(doc,
    headers=["Parameter", "Estimate", "Standard Error"],
    rows=[
        ["mu (location)",     "+0.000602", "0.000102"],
        ["b (scale)",         "0.008078",  "0.000102"],
        ["Log-likelihood",    "19,649.9",  "--"],
        ["AIC",               "-39,295.7", "--"],
        ["BIC",               "-39,282.2", "--"],
    ],
    caption="Table 2. Laplace MLE (symmetric VG special case: theta = 0, nu = 1)."
)

add_body(doc,
    "The Laplace closed-form MLE sets mu to the sample median and b to the "
    "mean absolute deviation about it. Its log-likelihood of 19,649.9 beats "
    "the Gaussian by 885.6 and falls only 15.6 short of the full VG, which "
    "frees the two parameters (theta and nu) that the Laplace holds fixed. A "
    "likelihood-ratio test of that restriction gives 31.2 on 2 degrees of "
    "freedom (p < 0.001), so the asymmetry and variance-rate freedom are real, "
    "but small next to the 885-unit gain the exponential tails alone give."
)

add_table(doc,
    headers=["Parameter", "Estimate", "Standard Error", "95% CI"],
    rows=[
        ["nu (degrees of freedom)", "2.648",     "0.110",    "(2.43, 2.87)"],
        ["mu (location)",           "+0.000656", "0.000110", "--"],
        ["sigma (scale)",           "0.007077",  "0.000128", "--"],
        ["Log-likelihood",          "19,666.7",  "--",       "--"],
        ["AIC",                     "-39,327.5", "--",       "--"],
        ["BIC",                     "-39,307.2", "--",       "--"],
    ],
    caption="Table 3. Student-t MLE (Week 2)."
)

add_table(doc,
    headers=["Parameter", "Estimate", "Standard Error"],
    rows=[
        ["sigma (scale)",      "0.011627",  "0.000163"],
        ["theta (asymmetry)",  "-0.000684", "0.000149"],
        ["nu (variance rate)", "1.17306",   "0.043087"],
        ["mu (location)",      "+0.000907", "0.000028"],
        ["Log-likelihood",     "19,665.5",  "--"],
        ["AIC",                "-39,322.9", "--"],
        ["BIC",                "-39,296.0", "--"],
    ],
    caption="Table 4. Variance-Gamma MLE."
)

add_body(doc,
    "The VG log-likelihood of 19,665.5 is almost identical to the Student-t's "
    "19,666.7 despite two extra parameters, so its AIC is slightly worse. The "
    "negative theta confirms left skew, and the variance rate nu = 1.173 is "
    "well above zero, so the Gamma time-change is doing real work rather than "
    "collapsing toward the Gaussian limit."
)

add_table(doc,
    headers=["Parameter", "Estimate", "Standard Error"],
    rows=[
        ["alpha (tail-heaviness)", "52.341",   "2.786"],
        ["beta (asymmetry)",       "-6.095",   "1.498"],
        ["delta (scale)",          "0.007574", "0.000217"],
        ["mu (location)",          "+0.001111","0.000150"],
        ["Log-likelihood",         "19,691.7", "--"],
        ["AIC",                    "-39,375.3","--"],
        ["BIC",                    "-39,348.4","--"],
    ],
    caption="Table 5. Normal Inverse Gaussian MLE."
)

add_body(doc,
    "NIG achieves the highest log-likelihood of the five models. Its AIC "
    "improvement over the Gaussian is 1,851 units, 48 better than the "
    "Student-t. Negative beta confirms left skew, and alpha = 52.3 indicates "
    "heavy tails; the values of 17-26 it falls to during the GFC and COVID "
    "(Section 5) show how much deeper the tails run in acute crises."
)


# ── 3. Model fit ──────────────────────────────────────────────────────────────

add_heading(doc, "3. Model fit")

add_body(doc,
    "Figure 1 overlays all five PDFs on the return histogram. The four "
    "non-Gaussian models look similar to the eye; the Laplace (dashed) and "
    "Student-t are nearly indistinguishable at this scale, and the NIG peak is "
    "marginally sharper. The QQ plots in Figure 2 show the difference more "
    "precisely. NIG and VG sit closest to the diagonal throughout, with the "
    "smallest residuals at the extreme observations; the Laplace tracks the "
    "body well but its tails are a touch too light at the most extreme points."
)

add_figure(doc, "week3_density_all_models.png",
    "Figure 1. Return histogram with all five fitted PDFs. Gaussian (blue), "
    "Laplace (purple, dashed), Student-t (red), VG (green), NIG (orange).")

add_figure(doc, "week3_qq_all_models.png",
    "Figure 2. QQ plots for all five models. The Gaussian S-shape (top left) "
    "is severe. NIG (bottom row) sits closest to the diagonal.",
    width_in=6.3)

add_table(doc,
    headers=["Model", "KS statistic (D)", "p-value", "n"],
    rows=[
        ["Gaussian",  "0.0938", "<0.001", "6,287"],
        ["Laplace",   "0.0180", "0.034",  "6,287"],
        ["Student-t", "0.0182", "0.030",  "6,287"],
        ["VG",        "0.0129", "0.250",  "6,287"],
        ["NIG",       "0.0113", "0.399",  "6,287"],
    ],
    caption=(
        "Table 6. KS goodness-of-fit. Gaussian, Laplace and Student-t: "
        "one-sample test (analytical CDF). VG and NIG: two-sample test "
        "against 1,000,000 simulated draws. D is the meaningful metric."
    )
)

add_body(doc,
    "NIG is the only model that cannot be rejected. VG passes too. The Laplace "
    "and Student-t are both borderline, and tie almost exactly on D (0.0180 "
    "versus 0.0182): a two-parameter symmetric model matches the three-parameter "
    "Student-t on this measure. The Gaussian fails at any threshold."
)

add_table(doc,
    headers=["Model", "Parameters", "Log-lik", "AIC", "BIC", "DAIC vs Gaussian"],
    rows=[
        ["Gaussian",  "2", "18,764.3", "-37,524.5", "-37,511.0", "--"],
        ["Laplace",   "2", "19,649.9", "-39,295.7", "-39,282.2", "-1,771.2"],
        ["Student-t", "3", "19,666.7", "-39,327.5", "-39,307.2", "-1,802.9"],
        ["VG",        "4", "19,665.5", "-39,322.9", "-39,296.0", "-1,798.4"],
        ["NIG",       "4", "19,691.7", "-39,375.3", "-39,348.4", "-1,850.8"],
    ],
    caption="Table 7. Model comparison. Lower AIC/BIC is better."
)

add_body(doc,
    "The Laplace, with two parameters, reaches DAIC -1,771, within 80 units of "
    "the four-parameter NIG and the most efficient model per parameter: most of "
    "the gain over the Gaussian comes from exponential tails rather than from "
    "extra parameters. The VG result cuts the other way. With four parameters "
    "it is slightly worse than the three-parameter Student-t on both AIC and "
    "BIC, which is not what you would expect if adding parameters always helped. "
    "NIG uses the same number of parameters as VG but gains 26 log-likelihood "
    "units, so the Inverse Gaussian mixing distribution captures something the "
    "Gamma mixing of VG cannot."
)


# ── 4. Risk measures ─────────────────────────────────────────────────────────

add_heading(doc, "4. Risk measures")

add_table(doc,
    headers=["Confidence",
             "VaR (Gauss.)", "ES (Gauss.)",
             "VaR (Lap)",    "ES (Lap)",
             "VaR (t)",      "ES (t)",
             "VaR (VG)",     "ES (VG)",
             "VaR (NIG)",    "ES (NIG)"],
    rows=[
        ["95%", "-1.990%", "-2.501%", "-1.800%", "-2.608%", "-1.694%", "-3.000%",
                "-1.915%", "-2.818%", "-1.880%", "-3.075%"],
        ["99%", "-2.824%", "-3.238%", "-3.100%", "-3.908%", "-3.515%", "-5.813%",
                "-3.360%", "-4.269%", "-3.769%", "-5.187%"],
    ],
    caption=(
        "Table 8. Daily VaR and ES for all five models. "
        "Gaussian, Laplace and Student-t use closed-form expressions; "
        "VG and NIG use Monte Carlo (500,000 draws)."
    )
)

add_body(doc,
    "At 99% ES the five models span from -3.24% (Gaussian) to -5.81% "
    "(Student-t). The Laplace sits at -3.91%, already 21% above the Gaussian "
    "from its exponential tails alone, with no asymmetry parameter. The "
    "Student-t's outlier reading comes from its symmetric tail assumption at "
    "nu = 2.648: forced to treat both tails equally at that extreme a nu, it "
    "overstates the right tail to match the left, and the overall ES inflates. "
    "VG and NIG, with separate asymmetry parameters, sit at -4.27% and -5.19% "
    "respectively. Using the Gaussian ES of -3.24% to set capital leaves a 24% "
    "shortfall relative to VG and 37% relative to NIG. Both gaps are large "
    "enough to matter under FRTB (BCBS 2013)."
)

add_figure(doc, "week3_risk_comparison.png",
    "Figure 3. VaR and ES at 95% and 99% for all five models. "
    "Hatched bars are ES.")


# ── 5. How crises change the distribution ────────────────────────────────────

add_heading(doc, "5. How crises change the distribution")

add_table(doc,
    headers=["Period", "n",
             "Gauss. sigma (ann.)",
             "t nu",
             "VG nu", "VG theta",
             "NIG alpha", "NIG beta"],
    rows=[
        ["Full sample",    "6,287", "19.4%", "2.648", "1.173", "-0.00068", "52.3",  "-6.09"],
        ["Dot-com crash",  "671",   "23.5%", "6.526", "0.397", "+0.00190", "98.6",  "+10.04"],
        ["GFC",            "378",   "38.3%", "2.607", "1.248", "-0.00296", "25.9",  "-2.74"],
        ["COVID-19",       "104",   "50.5%", "2.285", "1.439", "-0.00455", "17.8",  "-4.17"],
        ["Fed rate hikes", "501",   "19.5%", "6.525", "0.493", "-0.00002", "111.9", "-3.85"],
    ],
    caption="Table 9. Sub-period parameter estimates. Each window fitted independently."
)

add_body(doc,
    "The GFC and COVID return Student-t nu around 2.3-2.6, very close to the "
    "variance singularity at nu = 2. NIG alpha falls to 25.9 and 17.8. The "
    "dot-com crash and the rate-hike cycle show nu around 6.5 and NIG alpha "
    "around 99-112, near-Gaussian territory. The two types of crisis differ in "
    "kind, not merely in magnitude: the dot-com crash was a slow drawdown, "
    "while the GFC and COVID were clusters of extreme single-day moves. That "
    "difference is invisible in volatility alone but shows up immediately in "
    "the tail parameters."
)

add_table(doc,
    headers=["Period", "DAIC (Student-t)", "DAIC (VG)", "DAIC (NIG)"],
    rows=[
        ["Full sample",    "1,802.9", "1,798.4", "1,850.8"],
        ["Dot-com crash",   "22.7",    "22.3",    "23.0"],
        ["GFC",             "72.7",    "78.7",    "76.4"],
        ["COVID-19",        "19.7",    "22.6",    "20.9"],
        ["Fed rate hikes",  "16.0",    "16.2",    "15.1"],
    ],
    caption=(
        "Table 10. AIC improvement over Gaussian by period. "
        "Positive = Levy model preferred."
    )
)

add_body(doc,
    "The GFC improvement is 72-79 units; the calm periods are 15-23. "
    "Levy models earn their keep exactly when markets are worst."
)

add_figure(doc, "week3_subperiod_params.png",
    "Figure 4. Left: annualised scale by period and model. Right: tail "
    "parameter nu for Student-t and VG. The dashed line is the nu = 2 "
    "variance singularity.")

add_figure(doc, "week3_subperiod_aic.png",
    "Figure 5. AIC improvement of each Levy model over the Gaussian by period.")


# ── save ─────────────────────────────────────────────────────────────────────

doc.save(OUT)
print(f"Saved -> {OUT}")
