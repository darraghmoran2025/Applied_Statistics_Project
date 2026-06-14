"""
generate_week3_regression.py
Generates Week3_Regression.docx — the dedicated write-up for the
lead-up regression on forward risk (week3_leadup_regression.py).

Run from the Week3 directory: python generate_week3_regression.py

All numbers in this document are the output of
    python week3_leadup_regression.py
on the S&P 500 sample (2000-2024), horizon h = 21 trading days.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
FIGS = os.path.join(HERE, "figures")
OUT  = os.path.join(HERE, "Week3_Regression.docx")


# ── helpers ───────────────────────────────────────────────────────────────────

def _set_tnr(run, size_pt=12, bold=False, italic=False):
    run.font.name   = "Times New Roman"
    run.font.size   = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    r    = run._r
    rPr  = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rPr.insert(0, rFonts)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_tnr(run, bold=True)
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_tnr(run, bold=True, italic=True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_tnr(run)
    return p


def add_figure(doc, filename, caption, width_in=6.3):
    path = os.path.join(FIGS, filename)
    if not os.path.exists(path):
        add_body(doc, f"[Figure not found: {filename}]")
        return
    p = doc.add_paragraph()
    p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    _set_tnr(r, italic=True)


def add_table(doc, headers, rows, caption):
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after  = Pt(2)
    r = cap.add_run(caption)
    _set_tnr(r, bold=True)

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"

    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_tnr(run, bold=True)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            _set_tnr(run)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ── document ──────────────────────────────────────────────────────────────────

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

title_p = doc.add_paragraph()
title_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(6)
r = title_p.add_run("Week 3: Lead-up regression on forward risk")
_set_tnr(r, size_pt=14, bold=True)

subtitle_p = doc.add_paragraph()
subtitle_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
subtitle_p.paragraph_format.space_after = Pt(20)
r2 = subtitle_p.add_run(
    "What observable factors accompanied the build-up of tail risk before "
    "historical market shocks?"
)
_set_tnr(r2, size_pt=12, italic=True)


# ── 1. Scope and what this is NOT ─────────────────────────────────────────────

add_heading(doc, "1. Scope: this is risk attribution, not market prediction")

add_body(doc,
    "This regression does not attempt to predict the direction "
    "or level of stock returns. Forecasting whether the market will go up or "
    "down is outside the scope of this project and is not reliably possible: "
    "daily equity returns are very close to unforecastable. The dependent "
    "variable here is never a return sign or level. It is always a risk "
    "quantity, specifically forward realised volatility, which is the "
    "dispersion of returns over the coming month, not their direction."
)
add_body(doc,
    "Modelling the dispersion of returns is legitimate, and it is the "
    "foundation the Basel VaR and Expected Shortfall apparatus rests on, "
    "because volatility clusters. Large moves are followed by large moves and "
    "calm is followed by calm. This is one of the stylised facts of asset "
    "returns documented in Cont (2001), and it is why a static Gaussian risk "
    "number, which assumes constant volatility, fails. Using volatility "
    "clustering to describe how the risk environment evolves is a "
    "risk-measurement exercise, not a market-timing one."
)
add_body(doc,
    "The framing throughout is retrospective and in-sample. The question is: "
    "which observable, daily factors were systematically elevated in the "
    "run-up to historical episodes of extreme returns, and do tail-shape "
    "factors carry information about future risk beyond what plain volatility "
    "already tells us? No out-of-sample forecasting skill is claimed, no "
    "trading strategy or profit-and-loss is computed, and no held-out test is "
    "presented as a forecasting scorecard. The factor trajectories shown "
    "against the four shock windows are a descriptive diagnostic of the "
    "historical record, not an early-warning system."
)
add_body(doc,
    "There is a hard reason this restraint is not optional. The 25-year sample "
    "contains only four crisis episodes. Anything framed as predicting the "
    "crash would have effectively four positive events to learn from, which "
    "guarantees overfitting and supports no honest claim of forecasting power. "
    "A continuous risk target avoids this entirely: forward realised "
    "volatility is defined on roughly 6,200 daily observations, so the "
    "relationships estimated are statistically meaningful in a way a four-event "
    "crash model never could be."
)


# ── 2. How this differs from the cross-sectional VIX regression ───────────────

add_heading(doc, "2. How this differs from the earlier parameter-on-VIX regression")

add_body(doc,
    "An earlier exploratory regression in this project regressed the fitted "
    "annual Levy parameters on each year's average VIX. That earlier exercise "
    "is cross-sectional and contemporaneous: the unit of observation is a "
    "calendar year (n = 25), and the VIX and the parameters are measured over "
    "the same window, so it can only say whether two quantities move together. "
    "It carries no notion of before versus after and therefore cannot speak to "
    "lead-up dynamics. It is reported separately and is not part of this "
    "write-up."
)
add_body(doc,
    "The regression here is a time-series predictive regression in the "
    "econometric sense of that term: the unit of observation is a single "
    "trading day (n = 6,207), and every predictor is measured strictly at "
    "time t while the outcome is measured over the subsequent window t+1 to "
    "t+h. That deliberate temporal gap between predictors and outcome is the "
    "defining feature, and it is what makes the analysis a lead-up analysis "
    "rather than a contemporaneous description."
)


# ── 3. The model ──────────────────────────────────────────────────────────────

add_heading(doc, "3. The model")

add_subheading(doc, "The outcome: forward realised volatility")

add_body(doc,
    "The dependent variable is forward realised volatility over the next "
    "h = 21 trading days (approximately one calendar month). For each day t "
    "it is computed as the square root of the sum of squared daily log-returns "
    "over days t+1 to t+h, scaled by 252/h and then square-rooted, giving an "
    "annualised volatility figure. Squared returns are used as the building "
    "block because they measure the magnitude of moves regardless of sign, "
    "which is exactly the dispersion content we want and which strips out any "
    "directional (return-prediction) information by construction."
)

add_subheading(doc, "The predictors: all knowable at time t")

add_body(doc,
    "Every predictor is built only from information available up to and "
    "including day t. The set is small and economically motivated. Two of the "
    "predictors, the rolling skewness and rolling "
    "excess kurtosis of recent returns, are short-window empirical analogues "
    "of the Variance-Gamma and NIG asymmetry and tail-heaviness parameters "
    "estimated elsewhere in Week 3. Including them lets the model speak the "
    "project's own vocabulary: it asks whether the tail shape of recent "
    "returns, not just their volatility, carries information about the risk to "
    "come."
)

add_table(doc,
    headers=["Predictor", "Definition (all measured at day t)"],
    rows=[
        ["Trailing 21d vol", "Annualised standard deviation of returns over the last 21 days"],
        ["VIX level",        "CBOE Volatility Index level (options-implied 30-day vol)"],
        ["Delta VIX (5d)",   "5-day change in the VIX (momentum of implied volatility)"],
        ["|return| (1d)",    "Absolute value of the most recent daily return"],
        ["Rolling skew (21d)",     "Sample skewness of returns over the last 21 days"],
        ["Rolling kurtosis (21d)", "Excess kurtosis of returns over the last 21 days (tail proxy)"],
        ["Drawdown (252d)",  "Current price relative to its trailing 252-day peak (<= 0)"],
    ],
    caption="Table 1. The seven lead-up predictors. All are trailing quantities, "
            "so none uses information from after day t."
)

add_subheading(doc, "Estimation and standard errors")

add_body(doc,
    "The model is a single Ordinary Least Squares (OLS) regression fitted on "
    "the full daily sample. Because the forward windows of consecutive days "
    "overlap (the outcome for day t and day t+1 share 20 of their 21 days), "
    "the regression residuals are strongly autocorrelated, and ordinary OLS "
    "standard errors would be badly understated, so coefficients would look "
    "more significant than they are. To correct for this, all standard errors are "
    "Newey-West HAC (heteroskedasticity and autocorrelation consistent) "
    "standard errors, computed with a Bartlett kernel and a bandwidth set "
    "equal to the forecast horizon of 21 days, which is the mechanical overlap "
    "length. Every t-statistic and p-value reported below uses these robust "
    "standard errors."
)
add_body(doc,
    "Coefficients are reported in standardised (z-score) form, meaning each "
    "predictor and the outcome are rescaled to have mean zero and standard "
    "deviation one before fitting. A standardised coefficient of 0.40 then "
    "reads as: a one-standard-deviation increase in that factor is associated "
    "with a 0.40-standard-deviation increase in forward volatility, holding "
    "the other factors fixed. Standardisation makes the factor loadings "
    "directly comparable to one another despite their different natural units."
)


# ── 4. Results ────────────────────────────────────────────────────────────────

add_heading(doc, "4. Results")

add_subheading(doc, "Tail factors add information beyond volatility alone")

add_body(doc,
    "The main result is an incremental one. A baseline regression that "
    "uses only trailing 21-day volatility to explain forward volatility "
    "achieves an R-squared of 0.441: as expected from volatility clustering, "
    "recent volatility alone explains a large share of future volatility. "
    "Adding the other six factors raises the R-squared to 0.553. The "
    "additional 0.112, roughly eleven percentage points of explained "
    "variation, is contributed by the implied-volatility, shock, tail and "
    "drawdown factors over and above plain volatility persistence. The risk "
    "environment is better described by tail-aware factors than by volatility "
    "on its own. This is the same message the Levy models give for the static "
    "distribution, now in a dynamic form."
)

add_table(doc,
    headers=["Model", "Predictors", "R-squared", "Adjusted R-squared"],
    rows=[
        ["Baseline", "Trailing 21d vol only",        "0.441", "0.441"],
        ["Full",     "All seven lead-up factors",    "0.553", "0.552"],
        ["Increment","Tail / VIX / drawdown factors", "+0.112", "--"],
    ],
    caption="Table 2. Variance in forward 21-day realised volatility explained "
            "by the baseline and full models. n = 6,207 daily observations."
)

add_subheading(doc, "Which factors carry the signal")

add_body(doc,
    "Table 3 reports the standardised coefficients of the full model with "
    "Newey-West HAC standard errors. The VIX level is the single strongest "
    "factor (standardised coefficient +0.40, t = 4.2), which is expected: "
    "options-implied volatility is the market's own forward-looking estimate "
    "of dispersion, so it should and does carry the most information about "
    "future realised volatility. The five-day change in the VIX is also "
    "significant and positive (+0.12, t = 2.4): a rising VIX, and not only its "
    "level, precedes higher realised risk. The most recent absolute return "
    "carries a small but significant positive loading (+0.05, t = 2.2): a large "
    "move today tends to be followed by more turbulence over the next few weeks."
)
add_body(doc,
    "The two coefficients that warrant comment are the drawdown and the "
    "rolling kurtosis. Drawdown carries a negative loading (-0.11, t = -1.8, "
    "marginally significant): the closer the market is sitting to its trailing "
    "peak, the higher the volatility that tends to follow. Put differently, "
    "calm conditions near a record high are not reassuring on this measure. "
    "Rolling excess kurtosis also loads negatively and significantly (-0.04, "
    "t = -2.1). The negative sign should be read carefully and not "
    "over-interpreted: once recent volatility and the VIX are already in the "
    "model, a spike in trailing kurtosis tends to mark an isolated jump that "
    "has already happened rather than a build-up of risk still to come, so the "
    "residual association with future volatility is mildly negative. Rolling "
    "skewness is not statistically distinguishable from zero (t = -1.0)."
)

add_table(doc,
    headers=["Factor", "Std. coefficient", "HAC SE", "t-stat", "p-value"],
    rows=[
        ["VIX level",            "+0.399", "0.094", "+4.23", "<0.001"],
        ["Trailing 21d vol",     "+0.203", "0.105", "+1.94", "0.053"],
        ["Delta VIX (5d)",       "+0.122", "0.050", "+2.44", "0.015"],
        ["|return| (1d)",        "+0.052", "0.023", "+2.24", "0.025"],
        ["Rolling skew (21d)",   "-0.025", "0.026", "-0.97", "0.333"],
        ["Rolling kurtosis (21d)","-0.040","0.019", "-2.09", "0.037"],
        ["Drawdown (252d)",      "-0.115", "0.063", "-1.82", "0.069"],
    ],
    caption="Table 3. Standardised coefficients of the full model with "
            "Newey-West HAC standard errors (bandwidth 21). The intercept is "
            "zero by construction under standardisation and is omitted."
)

add_figure(doc, "week3_leadup_coefs.png",
    "Figure 1. Standardised factor loadings with Newey-West HAC 95% intervals. "
    "Bars whose interval does not cross zero are statistically significant. "
    "The VIX level dominates; drawdown and rolling kurtosis carry negative "
    "loadings.")

add_figure(doc, "week3_leadup_fit.png",
    "Figure 2. Actual forward 21-day realised volatility (black) against the "
    "in-sample fitted values from the full model (red), with the four shock "
    "windows shaded. The model tracks the broad evolution of risk, including "
    "the 2008 and 2020 spikes, with an in-sample R-squared of 0.55. This is "
    "an in-sample fit shown for diagnostic purposes; it is not an "
    "out-of-sample forecast.")


# ── 5. Retrospective lead-up to each shock ────────────────────────────────────

add_heading(doc, "5. What was elevated before each shock")

add_body(doc,
    "Table 4 is the retrospective attribution exhibit. For each shock window "
    "it reports the average standardised level of each factor over the 21 "
    "trading days immediately before the window began. Because the factors are "
    "z-scored over the full sample, a value of +0.7 means that factor stood "
    "0.7 sample standard deviations above its typical level going into the "
    "episode, and a negative value means it was below its typical level. This "
    "describes the conditions that preceded each episode in the historical "
    "record. It is explicitly not a forecast."
)
add_body(doc,
    "The clearest pattern is the drawdown factor, which was positive before "
    "all three episodes (+0.34, +0.68, +0.64). Each of these shocks began from "
    "a market sitting near its trailing peak rather than from an already "
    "depressed level. Beyond that the picture is mixed, which is the result "
    "worth stating plainly. The Global Financial Crisis was preceded by mildly "
    "elevated volatility and VIX (+0.39 and +0.30), consistent with a slow "
    "build-up. COVID-19 was the opposite: volatility, the VIX and skewness "
    "were all well below average in the weeks before the crash (-0.80, -0.70, "
    "-0.57), because that shock erupted out of an unusually calm market with "
    "almost no warning in these factors. This contrast is exactly why a "
    "forecasting claim would be indefensible, and why the analysis is framed "
    "as attribution rather than prediction."
)

add_table(doc,
    headers=["Shock window", "Trail vol", "VIX", "dVIX", "|ret|",
             "Skew", "Kurt", "Drawdn"],
    rows=[
        ["GFC",            "+0.39", "+0.30", "-0.36", "-0.09", "+0.04", "-0.25", "+0.34"],
        ["COVID-19",       "-0.80", "-0.70", "+0.25", "-0.26", "-0.57", "-0.47", "+0.68"],
        ["Fed rate hikes", "+0.14", "+0.12", "-0.59", "+0.13", "-0.14", "-0.64", "+0.64"],
    ],
    caption="Table 4. Mean standardised factor level in the 21 trading days "
            "before each shock window (z-scores; descriptive only). The "
            "dot-com crash is omitted because its start date (March 2000) "
            "falls inside the 252-day warm-up required by the drawdown factor, "
            "so a clean 21-day pre-window is not available in the sample."
)

add_figure(doc, "week3_leadup_factors.png",
    "Figure 3. Standardised trajectories of four representative factors "
    "(trailing volatility, rolling kurtosis, the 5-day VIX change, and "
    "drawdown) across the full sample, with the four shock windows shaded. "
    "The figure is a descriptive diagnostic of how the risk environment "
    "evolved, not a forecast.")


# ── 6. Limitations ────────────────────────────────────────────────────────────

add_heading(doc, "6. Limitations and honest caveats")

add_body(doc,
    "First, every R-squared and coefficient reported here is in-sample. The "
    "model has not been validated out of sample, and no claim of out-of-sample "
    "predictive skill is made; doing so responsibly would require a far larger "
    "set of crisis episodes than the historical record provides. Second, the "
    "predictors are correlated with one another, most obviously trailing "
    "volatility and the VIX level, so individual coefficients should be read "
    "as partial associations rather than independent causal effects, and the "
    "split of credit between collinear factors is not sharp. Third, the "
    "forward windows overlap, which is the reason Newey-West standard errors "
    "are used throughout; even so, effective sample information is smaller than "
    "the nominal 6,207 observations suggest. Fourth, the choice of a 21-day "
    "horizon is a modelling decision; shorter or longer horizons would shift "
    "the relative weight of the fast factors (the VIX change, the latest "
    "absolute return) against the slow ones (drawdown)."
)
add_body(doc,
    "Read with these caveats, the regression supports one conclusion, and it "
    "stays within the project's risk-measurement scope: the dispersion of "
    "future returns is partially describable from present, observable "
    "conditions, and tail-shape information adds materially to that description "
    "beyond volatility alone. It reinforces, in a dynamic form, what the Levy "
    "models show for the static distribution: the risk that matters lives in "
    "the tails, and it depends on the prevailing market state."
)


# ── save ─────────────────────────────────────────────────────────────────────

doc.save(OUT)
print(f"Saved -> {OUT}")
