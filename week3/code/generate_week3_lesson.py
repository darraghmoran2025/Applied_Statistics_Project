"""
generate_week3_lesson.py
Generates Week3_Lesson_v2.docx
Run from the Week3 directory: python generate_week3_lesson.py
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
FIGS = os.path.join(HERE, "figures")
OUT  = os.path.join(HERE, "Week3_Lesson.docx")


# ── helpers ───────────────────────────────────────────────────────────────────

def _set_tnr(run, size_pt=12, bold=False, italic=False):
    run.font.name   = "Times New Roman"
    run.font.size   = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    r    = run._r
    rPr  = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rPr.insert(0, rFonts)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_tnr(run, bold=True)
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_tnr(run, bold=True, italic=True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_tnr(run)
    return p


def add_figure(doc, filename, caption, width_in=6.3):
    path = os.path.join(FIGS, filename)
    if not os.path.exists(path):
        add_body(doc, f"[Figure not found: {filename}]")
        return
    p = doc.add_paragraph()
    p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    _set_tnr(r, italic=True)


def add_table(doc, headers, rows, caption):
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after  = Pt(2)
    r = cap.add_run(caption)
    _set_tnr(r, bold=True)

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"

    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_tnr(run, bold=True)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            _set_tnr(run)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ── document ──────────────────────────────────────────────────────────────────

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

title_p = doc.add_paragraph()
title_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(6)
r = title_p.add_run("Week 3: Levy process models for financial returns")
_set_tnr(r, size_pt=14, bold=True)

subtitle_p = doc.add_paragraph()
subtitle_p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
subtitle_p.paragraph_format.space_after = Pt(20)
r2 = subtitle_p.add_run(
    "A conceptual guide: the Laplace, Variance-Gamma (VG) and Normal Inverse "
    "Gaussian (NIG)"
)
_set_tnr(r2, size_pt=12, italic=True)


# ── 1. Where Week 2 left off ──────────────────────────────────────────────────

add_heading(doc, "1. Where Week 2 left off")

add_body(doc,
    "Week 2 fitted the Gaussian distribution and the Student-t distribution to "
    "6,287 daily log-returns on the S&P 500 (Standard and Poor's 500 Index, a "
    "market-capitalisation-weighted index of the 500 largest publicly traded "
    "companies in the United States) from January 2000 to December 2024. The "
    "Student-t gave nu-hat (the fitted degrees-of-freedom parameter, which "
    "controls how heavy the tails of the distribution are; smaller values mean "
    "fatter tails and a higher probability of extreme returns) = 2.648. The 99% "
    "Expected Shortfall (ES, the average loss on the worst 1% of days, the risk "
    "metric mandated by the Basel Committee on Banking Supervision under the "
    "Fundamental Review of the Trading Book, BCBS 2013, commonly abbreviated "
    "FRTB) was -5.81% under the Student-t against -3.24% under the Gaussian. "
    "A bank using the Gaussian to set capital under FRTB rules would hold 44% "
    "too little."
)
add_body(doc,
    "But the Student-t has a hard constraint it cannot get around: it is "
    "symmetric. Both tails must always be equal. Daily equity returns are not "
    "symmetric. Large losses happen more often than large gains of the same "
    "size. This is called negative skewness, the tendency for a distribution "
    "to have a longer or heavier left tail than right tail. The Student-t "
    "forces the tails to mirror each other regardless of the data. The two "
    "models introduced this week, the Variance-Gamma (VG) and the Normal "
    "Inverse Gaussian (NIG), remove that constraint. Each has a dedicated "
    "asymmetry parameter that lets the left tail be heavier than the right "
    "independently of how heavy the tails are overall."
)
add_body(doc,
    "There is a second issue with the Student-t. It ties all tail behaviour "
    "to a single parameter nu (degrees of freedom), which simultaneously "
    "controls how heavy the tails are and, because of the symmetry constraint, "
    "forces the same heaviness on both sides. VG and NIG separate these. Each "
    "has distinct parameters for scale (how spread out the distribution is), "
    "tail heaviness (how much probability sits in the extreme tails), and "
    "asymmetry (which tail is heavier). That separation makes their estimates "
    "economically interpretable in a way a single nu is not."
)


# ── 2. What a Levy process is ─────────────────────────────────────────────────

add_heading(doc, "2. What a Levy process is")

add_body(doc,
    "A Levy process is a continuous-time stochastic process, meaning a "
    "mathematical object that describes how something evolves randomly over "
    "time, with two defining properties. First, its increments are stationary: "
    "the distribution of changes over any interval of length t is the same "
    "regardless of when that interval starts. Second, its increments are "
    "independent: what happened in the past has no bearing on what happens "
    "next. Brownian motion, the continuous random process underlying the "
    "Gaussian model, is the simplest Levy process. VG and NIG are also Levy "
    "processes, but they allow for infinitely many small jumps per unit time, "
    "which is what produces heavier tails than Brownian motion alone."
)
add_body(doc,
    "Both VG and NIG belong to the class of Normal Variance-Mean Mixtures. "
    "The idea is that returns are still conditionally Gaussian (that is, if you "
    "knew the variance in advance, they would be normally distributed), but the "
    "variance and mean of that conditional Gaussian are themselves random "
    "variables drawn from a mixing distribution. Integrating out the mixing "
    "variable produces a heavy-tailed, potentially asymmetric marginal "
    "distribution. VG uses a Gamma distribution as its mixing distribution. "
    "NIG uses an Inverse Gaussian distribution. The Gamma distribution is a "
    "two-parameter family of continuous distributions on the positive real line, "
    "commonly used to model waiting times. The Inverse Gaussian distribution "
    "is the distribution of the first time a Brownian motion with drift reaches "
    "a fixed positive level; it also lives on the positive real line and has "
    "heavier tails than the Gamma. That extra tail weight in the mixing "
    "distribution is why NIG fits the S&P 500 returns better than VG. The "
    "mathematics works out to a closed-form probability density function (PDF), "
    "meaning an explicit formula that can be written down and evaluated "
    "numerically, making Maximum Likelihood Estimation (MLE), which is the "
    "process of finding the parameter values that make the observed data most "
    "probable, directly feasible."
)


# ── 3. The Laplace distribution ───────────────────────────────────────────────

add_heading(doc, "3. The Laplace distribution: the simplest heavy-tailed step")

add_body(doc,
    "Before the two Levy models, it is worth fitting the simplest distribution "
    "that improves on the Gaussian: the Laplace, also called the "
    "double-exponential. This was suggested by Neil, and it turns out to be one "
    "of the most informative single steps in the whole comparison. The Laplace "
    "density is f(x) = exp(-|x - mu| / b) / (2b), where mu is a location "
    "parameter and b is a scale parameter (both with the same roles as in the "
    "Gaussian). The defining feature is the absolute value |x - mu| in the "
    "exponent, in place of the Gaussian's squared term (x - mu)^2. That single "
    "change is what creates heavy tails: the Gaussian density dies off like "
    "exp(-x^2), extremely fast, whereas the Laplace dies off like exp(-|x|), "
    "much more slowly. The Laplace is literally two exponential distributions "
    "placed back to back, one decaying to the right of mu and one to the left."
)
add_body(doc,
    "The Laplace has exactly two parameters, the same count as the Gaussian, "
    "but its excess kurtosis (a measure of tail-heaviness; zero for the "
    "Gaussian) is exactly 3, meaning genuinely fatter tails. It is a fair "
    "like-for-like test of one idea in isolation: keep the model as simple as "
    "the Gaussian, change only the tail decay from quadratic to linear, and see "
    "how much of the gap to the data that closes. The answer (Section 7) is: "
    "most of it."
)

add_subheading(doc, "Why it belongs with the Levy models")

add_body(doc,
    "The Laplace is not just a convenient benchmark; it is a special case of "
    "the Variance-Gamma distribution introduced in the next section. The VG is "
    "built by running a Brownian motion on a random Gamma clock with variance "
    "rate nu and drift theta. When that clock has unit shape (which corresponds "
    "to nu = 1) the Gamma distribution becomes an Exponential distribution, and "
    "when the drift theta = 0 the construction is symmetric. Under those two "
    "restrictions, VG(mu, sigma, theta = 0, nu = 1) is exactly a Laplace "
    "distribution. So the Laplace is the symmetric, single-parameter-clock "
    "corner of the VG family. Fitting it first means that when the full VG is "
    "fitted afterwards, the extra log-likelihood it earns can be read directly "
    "as the value of freeing those two restrictions (asymmetry through theta, "
    "and a more variable clock through nu)."
)

add_subheading(doc, "Fitting it")

add_body(doc,
    "The Laplace is the one model here with a fully closed-form MLE (Maximum "
    "Likelihood Estimate), needing no numerical optimiser at all. The "
    "likelihood is maximised by setting mu to the sample median (not the mean) "
    "and b to the mean absolute deviation of the returns about that median. "
    "The median appears because minimising a sum of absolute deviations, which "
    "is what the |x - mu| term in the log-likelihood demands, is solved by the "
    "median, just as minimising squared deviations in the Gaussian case is "
    "solved by the mean. The fitted values on the full sample are mu = "
    "+0.000602 and b = 0.008078, with standard errors of 0.000102 on each "
    "(for the Laplace both standard errors equal b / sqrt(n))."
)


# ── 4. The Variance-Gamma distribution ────────────────────────────────────────

add_heading(doc, "4. The Variance-Gamma distribution")

add_subheading(doc, "The random-time intuition")

add_body(doc,
    "The cleanest way to understand VG is through the random-time construction "
    "from Madan, Carr and Chang (1998). Start with an ordinary Brownian motion "
    "with drift theta (the asymmetry parameter, which controls how much the "
    "process trends up or down between time-change events; negative theta "
    "produces left skew, meaning losses are more likely than gains of the "
    "same size) and volatility sigma (the scale parameter, which controls the "
    "width of the distribution contributed by the Brownian component alone). "
    "Now replace calendar time t with a random Gamma clock G(t)."
)
add_body(doc,
    "Instead of time passing at a steady rate, it passes faster on some days "
    "and slower on others. Think of it as market activity being uneven: on a "
    "quiet day, very little information arrives and prices barely move; on a "
    "busy day, news floods in and prices shift sharply. The Gamma clock G(t) "
    "captures this by having mean t, so on average the clock runs at the right "
    "speed, and variance nu times t, where nu is the variance rate of the Gamma "
    "time-change (the parameter that controls how much the clock fluctuates "
    "around its mean). Large nu means the clock is highly erratic, which "
    "produces heavy tails in the resulting return distribution because "
    "occasionally the clock runs very fast and prices make large moves. When "
    "nu = 0, the clock is deterministic and VG collapses to the Gaussian. The "
    "return over one period is X = theta times G plus sigma times W(G), where "
    "W is standard Brownian motion."
)

add_subheading(doc, "The four parameters")

add_body(doc,
    "mu is the location parameter. It shifts the whole distribution left or "
    "right without changing any other property. The full-sample MLE estimate "
    "is mu = +0.000907. This is the average daily log-return after the "
    "asymmetry and other parameters have been accounted for."
)
add_body(doc,
    "sigma is the scale parameter, the volatility of the underlying Brownian "
    "motion. The full-sample MLE estimate is sigma = 0.011627. This is smaller "
    "than the raw standard deviation of returns (0.012234) because sigma only "
    "captures the Brownian diffusion component; the random Gamma time-change "
    "adds additional spread on top."
)
add_body(doc,
    "theta is the asymmetry parameter, the drift of the underlying Brownian "
    "motion between events in the Gamma clock. Negative theta means the process "
    "trends slightly downward between clock ticks, producing a left-skewed "
    "return distribution. The full-sample MLE estimate theta = -0.000684 is "
    "negative and statistically significant, confirming that left skew is "
    "present in S&P 500 daily returns."
)
add_body(doc,
    "nu is the variance rate of the Gamma time-change. It controls tail "
    "heaviness by determining how erratic the clock is. Larger nu means the "
    "clock fluctuates more, which produces heavier tails. The full-sample MLE "
    "estimate nu = 1.173 means the Gamma clock has variance roughly equal to "
    "its mean, which is substantial clock randomness. In the Global Financial "
    "Crisis (GFC, the period from approximately October 2007 to March 2009 "
    "when the collapse of mortgage-backed securities triggered a near-failure "
    "of the global banking system) sub-period, nu = 1.248; in COVID-19 "
    "(the market shock of February to June 2020 caused by the global pandemic), "
    "nu = 1.439. Both are noticeably higher than in the dot-com crash "
    "(nu = 0.397) or the Federal Reserve rate-hike cycle (nu = 0.493), "
    "confirming that GFC and COVID produced fundamentally heavier-tailed "
    "return distributions."
)

add_subheading(doc, "The mathematics")

add_body(doc,
    "The VG PDF, derived by integrating out the Gamma mixing variable, is "
    "(Madan, Carr, Chang 1998, equation 2):"
)
add_body(doc,
    "f(x; mu, sigma, theta, nu) = "
    "2 * exp(theta*(x-mu)/sigma^2) / "
    "[sigma*sqrt(2*pi) * nu^(1/nu) * Gamma(1/nu)] "
    "* (|x-mu| / omega)^(1/nu - 1/2) * K_{1/nu - 1/2}(omega*|x-mu|/sigma^2)"
)
add_body(doc,
    "where omega = sqrt(2*sigma^2/nu + theta^2), Gamma(.) is the gamma "
    "function (a generalisation of the factorial to non-integer arguments), "
    "and K_v is the modified Bessel function of the second kind of order v. "
    "The Bessel function K_v is a standard mathematical function that appears "
    "naturally when solving certain differential equations; here it arises "
    "because the Gamma-Normal integral produces it analytically. For large "
    "|x-mu|, K_v(z) decays as sqrt(pi/(2z))*exp(-z), so the tails die "
    "exponentially but more slowly than the Gaussian's exp(-x^2) decay. This "
    "slower decay is what produces the heavier tails. Numerically, K_v(z) "
    "can underflow to floating-point zero for large z; the implementation "
    "uses the exponentially scaled version kve(v,z) = K_v(z)*exp(z) to "
    "prevent this."
)
add_body(doc,
    "The VG mean is mu + theta. The variance is sigma^2 + theta^2 * nu. The "
    "second term, theta^2 * nu, shows how the random time-change adds variance "
    "on top of the Brownian diffusion. Both terms contribute to the total "
    "spread; only the time-change term produces the heavy tails."
)


# ── 4. The Normal Inverse Gaussian distribution ───────────────────────────────

add_heading(doc, "5. The Normal Inverse Gaussian distribution")

add_subheading(doc, "The mixing idea")

add_body(doc,
    "NIG replaces the Gamma mixing distribution of VG with an Inverse "
    "Gaussian (IG) distribution. The IG distribution is the distribution of "
    "the first time a Brownian motion with drift reaches a fixed positive "
    "boundary. It lives on the positive real line, like the Gamma, but has "
    "heavier tails. Because the mixing distribution is heavier, the resulting "
    "return distribution is also heavier, which is why NIG fits the S&P 500 "
    "data better than VG."
)
add_body(doc,
    "The construction is: the return X given V follows a Normal distribution "
    "with mean mu + beta*V and variance V, where V is the random variance "
    "drawn from the IG mixing distribution. V has mean delta/gamma and "
    "variance delta/gamma^3. Here gamma = sqrt(alpha^2 - beta^2) is a derived "
    "quantity computed from the other parameters, not itself a free parameter. "
    "The mean of the conditional distribution is mu + beta*V. When beta is "
    "negative, large values of V push the mean downward and simultaneously "
    "widen the distribution, amplifying the left tail. This joint effect on "
    "mean and variance is what produces the asymmetry."
)

add_subheading(doc, "The four parameters")

add_body(doc,
    "mu is the location parameter. It shifts the distribution left or right. "
    "The full-sample MLE estimate mu = +0.001111 captures the average daily "
    "log-return after the asymmetry and tail parameters have been accounted for."
)
add_body(doc,
    "alpha is the tail-heaviness parameter. Larger alpha means lighter tails; "
    "as alpha approaches infinity the NIG distribution approaches the Gaussian. "
    "Smaller alpha means heavier tails. Values in the range 20 to 60 indicate "
    "genuinely heavy-tailed behaviour that is materially different from the "
    "Gaussian. The full-sample MLE estimate alpha = 52.3 sits in that range. "
    "During the GFC (alpha = 25.9) and COVID-19 (alpha = 17.8), alpha falls "
    "sharply, showing the distribution shifting deep into heavy-tail territory. "
    "During the dot-com crash (alpha = 98.6) and the Fed rate-hike cycle "
    "(alpha = 111.9), alpha is near the Gaussian regime."
)
add_body(doc,
    "beta is the asymmetry parameter. It must satisfy the constraint "
    "alpha > |beta|, meaning the absolute value of beta must be smaller than "
    "alpha, to ensure the distribution is mathematically proper. Negative beta "
    "produces a left-skewed distribution, consistent with equity returns. The "
    "full-sample MLE estimate beta = -6.095 is negative and statistically "
    "significant. During the dot-com crash, beta = +10.04, an unusual positive "
    "value discussed in Section 9."
)
add_body(doc,
    "delta is the scale parameter, analogous to sigma in the Gaussian. It "
    "controls the overall spread of the distribution. The full-sample MLE "
    "estimate delta = 0.007574 is close to the Student-t scale estimate of "
    "0.007077 from Week 2, which makes sense: both are scale parameters in "
    "heavy-tailed distributions fitted to the same data."
)

add_subheading(doc, "The mathematics")

add_body(doc,
    "The NIG PDF is:"
)
add_body(doc,
    "f(x; mu, alpha, beta, delta) = "
    "(alpha*delta / pi) * exp(delta*gamma + beta*(x-mu)) * "
    "K_1(alpha * sqrt(delta^2 + (x-mu)^2)) / sqrt(delta^2 + (x-mu)^2)"
)
add_body(doc,
    "where gamma = sqrt(alpha^2 - beta^2) and K_1 is the modified Bessel "
    "function of the second kind of order 1. The Bessel argument is "
    "alpha*sqrt(delta^2 + (x-mu)^2), which grows linearly in |x-mu| for "
    "large departures from the centre. This gives the same exponential tail "
    "decay as VG but with different curvature near the peak."
)
add_body(doc,
    "The NIG mean is mu + delta*beta/gamma. The variance is "
    "delta*alpha^2/gamma^3. The skewness (a measure of asymmetry; negative "
    "values indicate a longer left tail) is 3*beta / (alpha*sqrt(delta*gamma)), "
    "confirming that negative beta produces negative skewness."
)


# ── 5. Fitting by Maximum Likelihood ─────────────────────────────────────────

add_heading(doc, "6. Fitting by Maximum Likelihood Estimation (MLE)")

add_subheading(doc, "What MLE does")

add_body(doc,
    "Maximum Likelihood Estimation (MLE) finds the parameter values that make "
    "the observed data as probable as possible under the assumed distribution. "
    "Formally, for n observed returns r_1, ..., r_n, the MLE solves: "
    "maximise the log-likelihood L(theta) = sum of log f(r_i; theta) over all "
    "i, where f is the PDF of the distribution and theta is the vector of "
    "parameters. The log-likelihood is used instead of the raw likelihood "
    "because sums are easier to work with numerically than products of very "
    "small numbers."
)

add_subheading(doc, "Why numerical optimisation is needed")

add_body(doc,
    "For the Gaussian, MLE has closed-form solutions, meaning exact algebraic "
    "formulas: mu-hat (the MLE estimate of mu, the location parameter) equals "
    "the sample mean and sigma-hat (the MLE estimate of sigma, the scale "
    "parameter) equals the standard deviation of the data. The Laplace is also "
    "closed form, but with the median in place of the mean: mu-hat is the "
    "sample median and b-hat is the mean absolute deviation about it. VG and NIG "
    "have no such shortcuts. All four parameters must be optimised simultaneously over "
    "a four-dimensional parameter space. The implementation uses L-BFGS-B "
    "(Limited-memory Broyden-Fletcher-Goldfarb-Shanno with Bounds), a "
    "quasi-Newton numerical optimisation algorithm that approximates the "
    "curvature of the objective function from past gradient evaluations and "
    "supports box constraints, meaning upper and lower bounds on each "
    "parameter individually."
)
add_body(doc,
    "For VG, the constraints are sigma (scale) > 0 and nu (variance rate) > 0, "
    "both box constraints that L-BFGS-B handles directly. For NIG, the "
    "constraint alpha (tail-heaviness) > |beta| (absolute value of the "
    "asymmetry parameter) is a nonlinear constraint, meaning it cannot be "
    "expressed as a simple upper or lower bound on a single parameter. The "
    "solution is to reparametrise by defining xi = beta/alpha (the ratio of "
    "the asymmetry parameter to the tail-heaviness parameter). Whenever "
    "alpha > |beta|, xi lies in (-1, 1), so constraining xi to (-1, 1) is "
    "equivalent to enforcing the original constraint and is a box constraint "
    "that L-BFGS-B handles natively. The output converts xi back to beta "
    "using beta = xi * alpha."
)
add_body(doc,
    "Three starting points are tried for each model and the best result is "
    "kept. This guards against local optima, which are parameter values where "
    "the log-likelihood is higher than any nearby point but lower than the "
    "true global maximum. Standard errors are recovered from the observed "
    "Fisher Information matrix, which is the matrix of second derivatives of "
    "the negative log-likelihood at the MLE estimates. The Fisher Information "
    "matrix measures how sharply curved the likelihood surface is: a sharper "
    "peak means the parameters are more precisely identified and gives smaller "
    "standard errors. The Hessian (the matrix of second partial derivatives "
    "of the objective function with respect to the parameters) is approximated "
    "numerically using a four-point central-difference formula, and its "
    "inverse gives the asymptotic covariance matrix of the parameter estimates. "
    "The square roots of the diagonal entries of this matrix are the standard "
    "errors."
)

add_subheading(doc, "Risk measures by Monte Carlo simulation")

add_body(doc,
    "Value at Risk (VaR) is the loss threshold exceeded on the worst alpha% of "
    "days. At the 99% confidence level, VaR is the loss exceeded on only 1% of "
    "days. Expected Shortfall (ES) is the average loss on those worst days. ES "
    "is the more informative capital metric because it describes the inside of "
    "the tail, not just its edge. VaR and ES have no closed-form expressions "
    "under VG or NIG since neither distribution has a closed-form cumulative "
    "distribution function (the function giving the probability that a return "
    "falls below a given value). They are therefore computed by Monte Carlo "
    "simulation: draw a large sample from the fitted distribution and take "
    "empirical quantiles."
)
add_body(doc,
    "Monte Carlo simulation is a technique for approximating quantities that "
    "cannot be computed analytically by generating a large number of random "
    "draws from the relevant distribution and averaging or taking percentiles "
    "over those draws. With 500,000 draws per model the Monte Carlo error on "
    "VaR and ES is well below one basis point (one hundredth of one percent). "
    "For VG, sampling uses the Gamma-Normal mixture directly: draw G from a "
    "Gamma distribution with shape parameter 1/nu and scale parameter nu "
    "(where nu is the variance rate of the time-change), then draw the return "
    "X from a Normal distribution with mean mu + theta*G and variance "
    "sigma^2*G. For NIG, draw V from an Inverse Gaussian distribution with "
    "mean delta/gamma and scale delta^2, then draw X from Normal(mu + beta*V, V)."
)


# ── 6. What the data says ─────────────────────────────────────────────────────

add_heading(doc, "7. What the data says")

add_subheading(doc, "Model ranking")

add_body(doc,
    "NIG is the best-fitting model. Its log-likelihood, the sum of log-PDF "
    "values evaluated at the fitted parameters over all 6,287 observations, "
    "is 19,691.7. The Student-t scores 19,666.7, VG scores 19,665.5, and the "
    "two-parameter Laplace scores 19,649.9. The Akaike Information Criterion "
    "(AIC = 2k minus 2 times the log-likelihood, where k is the number of free "
    "parameters) penalises model complexity to prevent simply rewarding the "
    "model with the most parameters. AIC is lower for better models. The AIC "
    "improvement of NIG over the Gaussian is 1,851 units. The Bayesian "
    "Information Criterion (BIC = k times log(n) minus 2 times the "
    "log-likelihood, where n is the number of observations) imposes a heavier "
    "penalty for extra parameters than AIC; NIG also leads on BIC."
)
add_body(doc,
    "The Laplace result is the one to dwell on. With only two parameters, the "
    "same count as the Gaussian, it improves AIC by 1,771 units, which is 96% "
    "of the 1,851-unit improvement that the four-parameter NIG achieves. On a "
    "per-parameter basis it is the most efficient model in the comparison. The "
    "lesson is that the single most important defect of the Gaussian is not its "
    "symmetry or its precise tail shape, but simply that its tails are far too "
    "thin. Replacing the quadratic exponent with a linear one, at no extra "
    "parameter cost, recovers almost all of the available gain. The richer "
    "models then refine what is left: the full VG adds 15.6 log-likelihood "
    "units over the Laplace by freeing theta and nu, and NIG adds more again "
    "through its heavier mixing distribution."
)
add_body(doc,
    "The Kolmogorov-Smirnov (KS) goodness-of-fit test provides a formal "
    "statistical test of whether the fitted distribution matches the data. "
    "The KS statistic D measures the maximum absolute difference between the "
    "empirical cumulative distribution function of the data and the fitted "
    "model's cumulative distribution function. Smaller D means better fit. "
    "NIG gives D = 0.0113 with p-value 0.40, meaning there is a 40% probability "
    "of observing a D this large by chance if NIG is the true model. NIG cannot "
    "be rejected at any standard significance level. VG also passes (p = 0.25). "
    "The Student-t (p = 0.030) and the Laplace (D = 0.0180, p = 0.034) are both "
    "borderline, and they tie almost exactly on the KS statistic, so a "
    "two-parameter symmetric model matches the three-parameter Student-t on "
    "this measure. The Gaussian fails badly (D = 0.094, p < 0.001)."
)
add_body(doc,
    "VG with four parameters is slightly worse than the Student-t with three "
    "by AIC. This is a genuinely interesting result. The Gamma mixing "
    "distribution of VG captures a similar signal to the Student-t's "
    "nu (degrees-of-freedom) parameter but does not improve on it. NIG's "
    "Inverse Gaussian mixing distribution adds something extra because the IG "
    "is right-skewed; its variance V occasionally takes very large values, "
    "which simultaneously widens the conditional distribution and shifts its "
    "mean in the direction of beta. This joint effect amplifies the left tail "
    "of returns in a way the symmetric Gamma-based VG cannot replicate."
)

add_figure(doc, "week3_density_all_models.png",
    "Figure 1. Histogram of S&P 500 daily log-returns (January 2000 to "
    "December 2024) with all five fitted PDFs overlaid. Gaussian (blue), "
    "Laplace (purple, dashed), Student-t (red), VG (green), NIG (orange). The "
    "four non-Gaussian models are visually close to each other and all "
    "substantially better than the Gaussian at the centre and in the tails; "
    "the Laplace and Student-t are nearly coincident.")

add_figure(doc, "week3_qq_all_models.png",
    "Figure 2. Quantile-Quantile (QQ) plots for all five models. A QQ plot "
    "sorts the observed returns and plots each against the value the fitted "
    "distribution predicts at that rank. Points on the diagonal line indicate "
    "perfect fit. The Gaussian (top left) shows a severe S-shape, confirming "
    "fat tails. NIG (bottom row) sits closest to the diagonal, especially "
    "in the left tail where capital losses occur; the Laplace tracks the body "
    "well but its tails are slightly too light at the extremes.",
    width_in=6.3)

add_subheading(doc, "What the parameters mean economically")

add_body(doc,
    "The full-sample VG estimate of theta (the asymmetry parameter, the drift "
    "of the underlying Brownian motion) = -0.000684 confirms negative skewness "
    "in daily returns. The fitted nu (the variance rate of the Gamma "
    "time-change, controlling tail heaviness) = 1.173 means the random clock "
    "has variance roughly equal to its mean. For NIG, alpha (the "
    "tail-heaviness parameter; smaller means heavier tails) = 52.3 and beta "
    "(the asymmetry parameter; negative means left-skewed) = -6.095. Both are "
    "consistent with the stylised facts of equity returns documented in "
    "Cont (2001): negative skewness, excess kurtosis (fatter tails than a "
    "Gaussian), and tail decay that is slower than Gaussian but faster than a "
    "power law."
)


# ── 7. Risk measures ─────────────────────────────────────────────────────────

add_heading(doc, "8. Risk measures")

add_body(doc,
    "At 99% ES (the average loss on the worst 1% of days), the five models "
    "span a wide range. The Gaussian gives -3.24%, which is the most "
    "optimistic. The Laplace gives -3.91%, already 21% more conservative than "
    "the Gaussian from its exponential tails alone. VG gives -4.27%, NIG gives "
    "-5.19%, and the Student-t gives -5.81%. The Student-t's outlier reading is "
    "not because it is the best "
    "model; it is because at nu (degrees of freedom) = 2.648, very close to "
    "the boundary nu = 2 below which the variance becomes infinite, the "
    "Student-t forces extremely heavy symmetric tails. It overstates the right "
    "tail to match the left, and the resulting ES is inflated. VG and NIG, with "
    "their separate asymmetry parameters theta (VG) and beta (NIG), fit the "
    "two tails independently and produce more moderate ES values."
)
add_body(doc,
    "Using the Gaussian ES of -3.24% to set regulatory capital under FRTB "
    "(Fundamental Review of the Trading Book), the Basel Committee on Banking "
    "Supervision framework that replaced Value at Risk with Expected Shortfall "
    "as the primary capital metric, leaves a 24% shortfall relative to VG "
    "(-4.27%) and a 37% shortfall relative to NIG (-5.19%). This is the "
    "practical consequence of distributional misspecification."
)

add_figure(doc, "week3_risk_comparison.png",
    "Figure 3. VaR (Value at Risk, the loss threshold exceeded on the worst "
    "alpha% of days) and ES (Expected Shortfall, the average loss on those "
    "days) at 95% confidence (left panel) and 99% confidence (right panel) "
    "for all five models. Solid bars are VaR; hatched bars are ES. "
    "The ES bars grow faster than VaR as confidence rises because ES depends "
    "on the full shape of the tail, not just its edge.")

add_table(doc,
    headers=["Confidence",
             "ES (Gaussian)", "ES (Laplace)", "ES (Student-t)",
             "ES (VG)", "ES (NIG)"],
    rows=[
        ["95%", "-2.501%", "-2.608%", "-3.000%", "-2.818%", "-3.075%"],
        ["99%", "-3.238%", "-3.908%", "-5.813%", "-4.269%", "-5.187%"],
    ],
    caption=(
        "Table 1. Daily Expected Shortfall (ES) for all five models. "
        "ES is the average loss on days that exceed the VaR threshold. "
        "Gaussian, Laplace and Student-t computed analytically; VG and NIG "
        "computed by Monte Carlo simulation using 500,000 draws from the fitted "
        "distribution."
    )
)


# ── 8. The fingerprint of each crisis ────────────────────────────────────────

add_heading(doc, "9. The distributional fingerprint of each crisis")

add_body(doc,
    "Fitting all four models independently to each of the four market shock "
    "windows gives a distributional fingerprint of each period. A "
    "distributional fingerprint is simply the set of fitted parameter values "
    "for a given time window; comparing fingerprints across windows shows "
    "how the statistical character of returns changed from one crisis to another."
)
add_body(doc,
    "The GFC (Global Financial Crisis, approximately October 2007 to March "
    "2009) and COVID-19 windows return Student-t nu (degrees of freedom; "
    "values closer to 2 indicate fatter tails approaching infinite variance) "
    "of 2.607 and 2.285. VG nu (variance rate of the Gamma time-change, "
    "controlling tail heaviness; larger values mean heavier tails) rises to "
    "1.248 and 1.439. NIG alpha (tail-heaviness parameter; smaller values "
    "mean heavier tails) falls to 25.9 and 17.8. These are extreme readings "
    "across all three models simultaneously."
)
add_body(doc,
    "The dot-com crash (approximately March 2000 to October 2002, when the "
    "collapse of overvalued technology stocks led to a prolonged equity bear "
    "market) and the Federal Reserve rate-hike cycle (2022 to 2023, when "
    "rapid interest rate increases caused a broad repricing of growth assets) "
    "look completely different. Student-t nu = 6.526 and 6.525. NIG alpha "
    "= 98.6 and 111.9. Both periods are near the Gaussian regime in their "
    "tail structure. Gaussian annualised sigma (the standard deviation of "
    "daily log-returns scaled by the square root of 252 trading days per year, "
    "giving an annualised volatility figure) in the dot-com period is 23.5%, "
    "visibly elevated. But the tail parameters reveal that the elevation is "
    "in scale, not in tail shape. The GFC and COVID produced clustered extreme "
    "single-day moves. The dot-com crash was a gradual grinding decline. "
    "That difference is invisible in volatility but shows up immediately in "
    "nu, alpha, and beta."
)
add_body(doc,
    "One unusual result: VG theta (asymmetry parameter) = +0.00190 and NIG "
    "beta (asymmetry parameter) = +10.04 during the dot-com crash, both "
    "positive. Every other period shows negative asymmetry. The dot-com "
    "crash had more prolonged losses interspersed with volatile rallies, "
    "which a slightly right-skewed daily distribution can capture. The GFC "
    "and COVID had asymmetric, crash-dominated days where losses dramatically "
    "outnumbered gains of the same size."
)

add_figure(doc, "week3_subperiod_params.png",
    "Figure 4. Left: annualised scale parameter across all five periods for "
    "all four models. Right: tail parameter nu (degrees of freedom for "
    "Student-t; variance rate for VG) across the same periods. The dashed "
    "line at nu = 2 marks the variance singularity below which the Student-t "
    "variance becomes infinite. GFC and COVID approach it; all other periods "
    "stay well above.")

add_figure(doc, "week3_subperiod_aic.png",
    "Figure 5. AIC (Akaike Information Criterion) improvement of each Levy "
    "model over the Gaussian benchmark by sub-period. Positive values mean "
    "the Levy model is preferred. The GFC improvement of 72 to 79 units "
    "dwarfs the calmer periods (15 to 23 units), confirming that fat-tail "
    "models add the most value when markets are most severely dislocated.")


# Note: the regression analysis is reported separately in the dedicated
# regression write-up (generate_week3_regression.py → Week3_Regression.docx).


# ── 9. What comes next ────────────────────────────────────────────────────────

add_heading(doc, "10. What comes next")

add_body(doc,
    "The MLE (Maximum Likelihood Estimation) results reported throughout this "
    "document are point estimates: a single parameter vector that maximises "
    "the log-likelihood. Standard errors give an approximate measure of "
    "uncertainty based on the curvature of the likelihood surface, but they "
    "rely on asymptotic theory, meaning they become accurate as the number of "
    "observations n grows toward infinity. With n = 104 observations in the "
    "COVID-19 sub-period window, and with nu (variance rate, controlling tail "
    "heaviness) near the edge of the identifiable region where the likelihood "
    "surface becomes flat, the asymptotic approximation is not reliable for "
    "those sub-period estimates."
)
add_body(doc,
    "Weeks 4 and 5 address this using Bayesian estimation via NUTS "
    "(No-U-Turn Sampler, an efficient Markov Chain Monte Carlo algorithm that "
    "automatically tunes its step sizes and avoids the inefficiency of random "
    "walks by using gradient information to explore the parameter space) "
    "implemented in PyMC (a Python library for probabilistic programming that "
    "builds and samples from Bayesian models). Markov Chain Monte Carlo "
    "(MCMC) is a class of algorithms that generates a sequence of parameter "
    "values whose distribution converges to the posterior distribution, "
    "meaning the probability distribution over parameter values given the "
    "observed data. Instead of a single point estimate like nu = 1.173, NUTS "
    "produces thousands of samples from the full posterior distribution over "
    "nu, which can be used to compute probabilities like: what is the "
    "probability that nu exceeds 1.5 given this data? That is the question a "
    "risk manager actually needs answered when trying to judge whether a tail "
    "regime shift has occurred, and it cannot be answered with a point "
    "estimate alone."
)


# ── save ─────────────────────────────────────────────────────────────────────

doc.save(OUT)
print(f"Saved -> {OUT}")
